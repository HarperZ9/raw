"""Generate RAW Technical Handoff Document as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Styles setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# Code style
code_style = doc.styles.add_style('CodeBlock', 1)  # paragraph style
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(9)
code_font.color.rgb = RGBColor(0x20, 0x20, 0x20)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.left_indent = Inches(0.3)


def add_code_block(text):
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='CodeBlock')


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()  # spacer


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RAW (Rendering Advancement Workshop)')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Technical Handoff Document')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ('Author', 'Zain D. Harper (papacr0w)'),
    ('Date', '2026-03-20'),
    ('Recipient', 'Pascal Gilcher (MartyMcModding)'),
    ('Codebase', r'C:\Users\Zain\SKSE\Playground\RAW'),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run(value)
    run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Architecture',
    '   2.1 Two-DLL Communication',
    '   2.2 Render Phase Detection',
    '   2.3 Mid-Frame Dispatch Pipeline',
    '   2.4 Effect Registration Pattern',
    '   2.5 Depth Pipeline',
    '3. Complete File Inventory',
    '   3.1 Proxy',
    '   3.2 Core Infrastructure',
    '   3.3 Effect Renderers',
    '   3.4 HLSL Shaders',
    '4. The Core Problem: Depth Buffer Access',
    '   4.1 The Fundamental Challenge',
    '   4.2 What Was Tried',
    '   4.3 The Phase Detector Timing Problem',
    '   4.4 D3D11 SRV/DSV Hazard',
    '5. Issues Encountered (Chronological)',
    '6. Current State (2026-03-20)',
    '7. Options Moving Forward',
    '8. Recommended Path',
    '9. Code Quality Notes',
    '10. Build & Deploy',
    '11. Key Files to Start With',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('   '):
        for r in p.runs:
            r.bold = True

doc.add_page_break()

# ============================================================
# 1. PROJECT OVERVIEW
# ============================================================
doc.add_heading('1. Project Overview', level=1)

doc.add_paragraph(
    'RAW is a complete D3D11 rendering platform for Skyrim SE that aims to replace both '
    'ENB and Community Shaders. It consists of two DLLs:'
)

doc.add_paragraph(
    'd3d11.dll (proxy) \u2014 Replaces Skyrim\'s d3d11.dll via DLL search order. Wraps '
    'ID3D11Device, ID3D11DeviceContext, and IDXGISwapChain. Classifies the game\'s '
    'rendering into 9 phases via heuristics. Provides a ProxyInterface struct for '
    'cross-DLL communication.',
    style='List Bullet'
)

doc.add_paragraph(
    'RAW.dll (SKSE plugin) \u2014 Loaded by SKSE. Contains 28 effect renderers, all '
    'shaders rewritten from published papers. Dispatches effects mid-frame at 5 pipeline '
    'stages via callbacks from the proxy.',
    style='List Bullet'
)

p = doc.add_paragraph()
run = p.add_run('Scale: ')
run.bold = True
p.add_run('~54,600 lines of C++ across 89 source files, 84 external HLSL shaders.')

# ============================================================
# 2. ARCHITECTURE
# ============================================================
doc.add_heading('2. Architecture', level=1)

doc.add_heading('2.1 Two-DLL Communication', level=2)

doc.add_paragraph(
    'The proxy DLL intercepts D3D11CreateDeviceAndSwapChain, wraps the returned Device, '
    'Context, and SwapChain objects, and stores the real (unwrapped) pointers in a '
    'ProxyInterface singleton.'
)

doc.add_paragraph(
    'The SKSE plugin (RAW.dll) discovers the proxy at runtime via '
    'GetModuleHandle("d3d11.dll") + GetProcAddress("PG_GetProxyInterface"). It retrieves '
    'the real Device/Context/SwapChain pointers and registers callbacks for phase changes, '
    'PrePresent, and resize events.'
)

doc.add_paragraph(
    'The ProxyInterface (ProxyAPI.h:45-148) is a plain C struct with function pointers for '
    'callback registration, raw D3D11 pointers, HDR state, material pipeline stats, and '
    'G-buffer SRVs. No vtable, no COM \u2014 just a flat struct shared between the two DLLs.'
)

doc.add_heading('2.2 Render Phase Detection', level=2)

doc.add_paragraph(
    'RenderPhaseDetector (RenderPhaseDetector.cpp) classifies the game\'s current '
    'rendering phase by observing:'
)

doc.add_paragraph('RT changes (OnRTChange): Depth-only = DepthPrepass/ShadowMap. Full-res color+depth = GeometryMain. Color without depth = PostProcess.', style='List Bullet')
doc.add_paragraph('Draw calls (OnDraw): Draw(3, 0) non-indexed during GeometryMain with >=200 geometry draws = PostProcess transition.', style='List Bullet')
doc.add_paragraph('Shader binds (OnShaderBind): Known sky shader hashes = Sky phase.', style='List Bullet')
doc.add_paragraph('Viewport changes (OnViewportChange): Sub-backbuffer viewports during depth-only = ShadowMap.', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('9 Phases: ')
run.bold = True
p.add_run('Unknown \u2192 DepthPrepass \u2192 ShadowMap \u2192 GeometryMain \u2192 Decals \u2192 Sky \u2192 AlphaBlend \u2192 PostProcess \u2192 UI')

doc.add_heading('2.3 Mid-Frame Dispatch Pipeline', level=2)

doc.add_paragraph(
    'When RenderPhaseDetector fires a phase transition, callbacks propagate to '
    'PhaseDispatcher (PhaseDispatcher.cpp), which:'
)

steps = [
    'Maps the transition to a PipelineStage (PostDepthPrepass, PostGeometry, PostSky, PreUI, PrePresent)',
    'Calls SceneMatrices::UpdateFromNiCamera() for live camera data',
    'Saves full D3D11 state via D3D11StateBackup',
    'Unbinds all OM targets (prevents SRV/DSV read-write hazards)',
    'Calls RenderPipeline::ExecuteStage() which runs all registered passes at that stage',
    'Restores D3D11 state',
    'Invalidates the proxy\'s state redundancy cache',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('2.4 Effect Registration Pattern', level=2)

doc.add_paragraph('Each effect renderer is a singleton that:')
doc.add_paragraph('Initialize(dev, ctx, sc) \u2014 compiles shaders, creates GPU resources', style='List Bullet')
doc.add_paragraph('Registers with RenderPipeline::AddPass() at a specific stage + priority', style='List Bullet')
doc.add_paragraph('Execute(PassContext&) \u2014 dispatches compute or pixel shaders', style='List Bullet')
doc.add_paragraph('Exposes output SRV for the compositor or SRV injection', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Priority order at PostGeometry: ')
run.bold = True
p.add_run('GTAO (15) \u2192 ContactShadow (16) \u2192 SSGI (20) \u2192 SSR (25) \u2192 Skylighting (30) \u2192 SceneCompositor (90)')

doc.add_heading('2.5 Depth Pipeline', level=2)

add_code_block("""Game renders depth (D24_UNORM_S8_UINT, standard: near=0, far=1)
  -> HiZPyramid::BuildPyramid() runs at Present time
    -> CSCopy shader: output = 1.0 - input (standard -> reversed-Z)
    -> Downsample: 2x2 max for each mip level
    -> Result: R32_FLOAT pyramid at t19
  -> SharedGPUResources::Update() runs after HiZ
    -> Linearization: viewZ = N*F / (N + z*(F-N)) using reversed-Z
    -> Result: R32_FLOAT linear depth at t31""")

p = doc.add_paragraph()
run = p.add_run('Key constants: ')
run.bold = True
p.add_run('near=15 (~21cm), far=353840 (~5km), 1 unit \u2248 1.43cm')

# ============================================================
# 3. FILE INVENTORY
# ============================================================
doc.add_heading('3. Complete File Inventory', level=1)

doc.add_heading('3.1 Proxy (src/d3d11_proxy/ \u2014 26 files, ~5,800 LOC)', level=2)

add_table(
    ['File', 'LOC', 'Purpose'],
    [
        ['proxy_main.cpp', '683', 'DLL entry, LazyInit, ProxyInterface, D3D11CreateDevice* exports'],
        ['WrappedDevice.cpp/h', '512', 'ID3D11Device wrapper (mostly passthrough)'],
        ['WrappedContext.cpp/h', '1180', 'ID3D11DeviceContext wrapper \u2014 draw/RT/shader hooks, state cache'],
        ['WrappedSwapChain.cpp/h', '1014', 'IDXGISwapChain wrapper \u2014 Present hook, depth capture, resize'],
        ['RenderPhaseDetector.cpp/h', '460', '9-phase heuristic classifier'],
        ['MaterialPipeline.cpp/h', '1074', 'DXBC bytecode patching for G-buffer extraction'],
        ['AlbedoExtractor.cpp/h', '762', 'Albedo G-buffer extraction via shader patching'],
        ['ShaderManager.cpp/h', '323', 'Shader hash tracking'],
        ['CBDirtyTracker.cpp/h', '228', 'Constant buffer upload optimization'],
        ['OcclusionCuller.cpp/h', '575', 'GPU occlusion culling'],
        ['DepthOwnership.cpp/h', '231', 'Attempted depth buffer ownership (DISABLED)'],
        ['ProxyAPI.h', '148', 'Cross-DLL interface struct definition'],
        ['ProxyLog.h', '57', 'Proxy-side file logging'],
    ]
)

doc.add_heading('3.2 Core Infrastructure (src/core/ \u2014 30 infra files, ~8,500 LOC)', level=2)

add_table(
    ['File', 'LOC', 'Purpose'],
    [
        ['main.cpp', '680', 'Plugin entry, init order, frame update loop'],
        ['D3D11Hook.cpp/h', '1093', 'Proxy connection, Present callback, ImGui setup'],
        ['D3D11StateBackup.h', '154', 'Full D3D11 pipeline state save/restore'],
        ['PhaseDispatcher.cpp/h', '326', 'Maps proxy phase transitions to pipeline stages'],
        ['RenderPipeline.cpp/h', '417', 'Pass orchestration, stage execution'],
        ['RenderPassManager.cpp/h', '552', 'Fullscreen VS+PS infrastructure'],
        ['ComputeManager.cpp/h', '506', 'CS dispatch with state save/restore'],
        ['SceneData.cpp/h', '372', 'Camera matrices from NiCamera'],
        ['SharedGPUResources.cpp/h', '402', 'Linear depth, blue noise, vanilla params CB'],
        ['HiZPyramid.cpp/h', '388', 'Hierarchical depth with standard to reversed-Z conversion'],
        ['ClusteredLighting.cpp/h', '701', 'Forward+ (2048 lights, 16x16x32 clusters)'],
        ['MaterialClassifier.cpp/h', '210', 'Per-pixel material ID'],
        ['LuminanceHistogram.cpp/h', '696', '256-bin GPU histogram'],
        ['SRVInjector.cpp/h', '88', 'Shared SRV slot management'],
        ['ShaderLoader.cpp/h', '263', 'Hot-reload HLSL from disk'],
        ['ShaderCache.cpp/h', '246', 'FNV-1a bytecode disk cache'],
        ['ConfigManager.cpp/h', '329', 'INI persistence + 5 presets'],
        ['GPUResource.h', '75', 'CreateGPUTexture/CreateCB/UploadCB/SafeRelease helpers'],
        ['DXBCPatcher.cpp/h', '734', 'DXBC bytecode patching (MRT output injection)'],
        ['Debug/Diag (various)', '~2500', 'DebugGUI, ShaderDebug, BootDiag, GPUProfiler, FrameCapture'],
    ]
)

doc.add_heading('3.3 Effect Renderers (src/core/ \u2014 28 renderers, ~24,000 LOC)', level=2)

add_table(
    ['Renderer', 'LOC', 'Stage', 'SRV Slot', 'Status'],
    [
        ['GTAORenderer', '1047', 'PostGeometry:15', 't20', 'Compiles, depth issue blocks testing'],
        ['ContactShadowRenderer', '845', 'PostGeometry:16', 't28', 'Compiles, untested'],
        ['SSGIRenderer', '1625', 'PostGeometry:20', 't26', 'Compiles, untested'],
        ['SSRRenderer', '1084', 'PostGeometry:25', 't27', 'Compiles, untested'],
        ['SkylightingRenderer', '1098', 'PostGeometry:30', 't29', 'Compiles, untested'],
        ['SceneCompositor', '587', 'PostGeometry:90', '\u2014', 'Compiles, compositor pipeline works'],
        ['BloomRenderer', '1203', 'PostUI', '\u2014', 'Compiles, untested'],
        ['DoFRenderer', '1642', 'PostUI', '\u2014', 'Compiles, untested'],
        ['ColorPipeline', '922', 'PostUI', '\u2014', 'Compiles, untested'],
        ['ToneMapManager', '824', 'PostUI', '\u2014', '12 operators implemented'],
        ['LensRenderer', '1306', 'PostUI', '\u2014', 'Compiles, untested'],
        ['UnderwaterRenderer', '1261', 'PostUI', '\u2014', 'Compiles, untested'],
        ['AtmosphereRenderer', '1022', '\u2014', '\u2014', 'LUT generation works'],
        ['VolumetricClouds', '1225', '\u2014', '\u2014', 'Compiles, deferred init'],
        ['FrameGenerator', '853', 'PrePresent', '\u2014', 'Compiles, deferred init'],
        ['TemporalSuperRes', '1088', 'PrePresent', '\u2014', 'Compiles, deferred init'],
        ['+ 10 more', '~4000', 'Various', 'Various', 'Compiles, deferred init'],
    ]
)

doc.add_heading('3.4 HLSL Shaders (84 files in Shaders/)', level=2)

doc.add_paragraph(
    'All 84 shaders are external .hlsl files hot-reloadable via F12. They cover: '
    'AO (3), contact shadows (2), skylighting (5), SSR (3), SSGI (4), bloom (5), '
    'DoF (5), lens (5), tone mapping (2), atmosphere (4), clouds (4), volumetric (2), '
    'clustered lighting (2), particle lights (3), underwater (4), denoising (5), '
    'color grading (1), HiZ (2), motion vectors (1), frame gen (2), TAA/TSR (3), '
    'material (3), debug (2), utility (2).'
)

# ============================================================
# 4. CORE PROBLEM
# ============================================================
doc.add_heading('4. The Core Problem: Depth Buffer Access', level=1)

p = doc.add_paragraph()
run = p.add_run(
    'This is the single blocking issue that has consumed the majority of development '
    'time and prevents any effect from producing correct output.'
)
run.bold = True

doc.add_heading('4.1 The Fundamental Challenge', level=2)

doc.add_paragraph(
    'Skyrim\'s depth buffer (D24_UNORM_S8_UINT) is created by the game engine without '
    'D3D11_BIND_SHADER_RESOURCE. This means:'
)

doc.add_paragraph('You cannot create an SRV directly on the game\'s depth texture', style='List Number')
doc.add_paragraph('You must either: (a) copy the depth to a shader-readable texture, or (b) create your own depth texture and substitute it', style='List Number')

doc.add_paragraph('Both approaches have been attempted. Both have problems.')

doc.add_heading('4.2 What Was Tried', level=2)

p = doc.add_paragraph()
run = p.add_run('Approach 1: Pre-Clear Depth Copy (WrappedSwapChain.cpp)')
run.bold = True

doc.add_paragraph(
    'Concept: Intercept ClearDepthStencilView in WrappedContext. Before the game clears '
    'its depth, copy the full depth texture to our own R24G8_TYPELESS texture with an SRV.'
)
doc.add_paragraph(
    'Result: Works for Present-time capture (complete depth from previous frame). But for '
    'mid-frame effects that need current-frame depth, the copy timing is wrong \u2014 the copy '
    'happens at the clear before geometry, not after geometry.'
)

p = doc.add_paragraph()
run = p.add_run('Approach 2: Depth Ownership (DepthOwnership.cpp)')
run.bold = True

doc.add_paragraph(
    'Concept: Create our own depth texture (R24G8_TYPELESS with both BIND_DEPTH_STENCIL '
    'and BIND_SHADER_RESOURCE). Intercept OMSetRenderTargets and substitute our DSV for '
    'the game\'s DSV. The game writes depth to our texture, and we can read it via SRV.'
)
doc.add_paragraph(
    'Result: DISABLED. The game wrote depth values to our texture, but the values were wrong '
    '\u2014 DepthProbe showed uniform 0.002695 instead of varying scene depth. Possible causes: '
    'the game\'s depth stencil state interacts poorly with a different texture; there are '
    'additional code paths binding the DSV that we weren\'t intercepting; format mismatch.'
)

p = doc.add_paragraph()
run = p.add_run('Approach 3: Present-Time HiZ Build (CURRENT STATE)')
run.bold = True

doc.add_paragraph(
    'Concept: Build HiZ and linearized depth only at Present time, where the game\'s depth '
    'buffer is guaranteed complete (all geometry drawn). Effects at PostGeometry read the '
    'previous frame\'s depth (one frame old).'
)
doc.add_paragraph(
    'Result: This is the current deployed state. One-frame-old depth should be invisible at '
    '60fps. However, it has not been confirmed working at all camera angles. The last test '
    'showed depth working only when looking at the ground (close objects) but failing at '
    'other angles.'
)

doc.add_heading('4.3 The Phase Detector Timing Problem', level=2)

doc.add_paragraph(
    'Independent of depth buffer access, the RenderPhaseDetector fires the '
    'GeometryMain \u2192 PostProcess transition too early. Skyrim has mid-geometry fullscreen passes:'
)

add_code_block("""Game render order:
1. Clear depth
2. Draw close geometry (DepthPrepass + early GeometryMain)
3. Mid-geometry fullscreen pass: Draw(3, 0)  <-- DETECTOR FIRES HERE
4. Draw distant geometry (rest of GeometryMain)
5. Sky, PostProcess, UI, Present""")

doc.add_paragraph(
    'The Draw(3, 0) heuristic (line 217 of RenderPhaseDetector.cpp) with a threshold of '
    '200 geometry draws triggers PostProcess before all geometry is drawn. This means: '
    'PostGeometry effects fire with incomplete scene depth; close objects have depth, '
    'distant objects don\'t; effects produce correct output when looking at the ground '
    '(only close geometry) but fail otherwise.'
)
doc.add_paragraph('Threshold was raised from 10 \u2192 200 but still fires too early for complex scenes.')

doc.add_heading('4.4 D3D11 SRV/DSV Hazard', level=2)

doc.add_paragraph(
    'D3D11 has a strict rule: you cannot simultaneously read (SRV) and write (DSV/RTV) to '
    'the same resource. When you bind a depth texture as both DSV and SRV, the SRV reads '
    'return zero silently \u2014 no error, no warning, just black.'
)
doc.add_paragraph(
    'The PhaseDispatcher unbinds all OM targets before executing effects (PhaseDispatcher.cpp '
    'line 196-200), which should prevent this. But the HiZ build at Present time also needs '
    'to unbind \u2014 HiZPyramid::BuildPyramid has an explicit PS SRV unbind loop for slots '
    '0-19 before dispatching.'
)

# ============================================================
# 5. ISSUES ENCOUNTERED
# ============================================================
doc.add_heading('5. Issues Encountered (Chronological)', level=1)

issues = [
    ('5.1 Reversed GTAO Bitmask (Phase 1)',
     'The original GTAO implementation had a reversed bitmask watermark from Marty\'s '
     'implementation. Identified in Marty\'s code review and was one of the reasons for '
     'the clean-room rewrite.'),
    ('5.2 Code Bloat (Phase 1)',
     'The original codebase had massive code bloat with non-functional effects. Marty\'s '
     'feedback: too much code, most of it doesn\'t work, SSGI was completely non-functional. '
     'Led to the clean-room restart.'),
    ('5.3 Depth Convention Confusion (Phase 2-3)',
     'Skyrim uses standard depth (near=0, far=1). The codebase went through multiple cycles '
     'of converting conventions, debugging, flip-flopping, then realizing the convention was '
     'correct all along and the real problem was elsewhere. Current state: HiZ converts '
     'standard\u2192reversed (1.0 - z). All effect shaders use reversed-Z conventions.'),
    ('5.4 Standard Depth Precision Issues (Phase 3)',
     'With standard depth, near=15, far=353840 (ratio 24,776:1), scene geometry falls in '
     'D24 range 0.001-0.01. This crushes precision. Normal reconstruction from depth '
     'gradients produced garbage \u2014 manifesting as kaleidoscope/broken glass patterns. '
     'Reversed-Z (which HiZ provides) maps the same geometry to 0.99-0.999.'),
    ('5.5 Phase Detector Premature Firing (Phase 3-4 \u2014 UNSOLVED)',
     'The Draw(3, 0) heuristic fires before all geometry is drawn. Root cause of most '
     'visible artifacts. Threshold raised from 10 \u2192 200 but still fires too early. '
     'Alternative detection methods (RT-change, shader hash) considered but not implemented.'),
    ('5.6 SRV/DSV Read-Write Hazard (Phase 3)',
     'Binding depth texture as both DSV and SRV causes silent zero reads. Solution '
     '(unbinding OM targets before dispatch) is implemented but was a source of '
     'black-screen bugs during development.'),
    ('5.7 RT Format Mismatch (Phase 3)',
     'Phase detector sometimes fires when a non-scene RT is bound (e.g., R8G8_UNORM temp '
     'textures). SceneCompositor has an RT format guard that rejects non-color formats, '
     'discovered after debugging black-screen/smearing artifacts.'),
    ('5.8 Blue Tint from AO (Phase 3)',
     'GTAO producing ao\u22480 (full occlusion) for all pixels, uniformly darkening the scene '
     'and letting Skyrim\'s blue ambient color show through. Caused by depth being '
     'zero/wrong at PostGeometry time.'),
    ('5.9 PreUI Dispatch Attempt (Phase 3)',
     'Effects moved to PreUI stage hoping for complete depth. The depth was correct, but '
     'the game\'s RTV at PreUI time is NOT the displayed scene RT \u2014 it\'s a different '
     'texture. Compositor output was invisible. Reverted to PostGeometry.'),
    ('5.10 Depth Ownership Wrong Values (Phase 3)',
     'Our substituted depth texture received uniform 0.002695 values instead of varying '
     'depth. Even after adding DSV substitution to both OMSetRenderTargets AND '
     'OMSetRenderTargetsAndUnorderedAccessViews. Root cause unclear \u2014 disabled and abandoned.'),
    ('5.11 Legacy/Ownership System Conflict (Phase 3)',
     'PreClearDepthCopy (legacy) and DepthOwnership both wrote to pi->gameDepthSRV, '
     'overwriting each other. Fixed by disabling legacy when ownership was active, but '
     'then ownership was disabled entirely.'),
    ('5.12 FP Depth Threshold Catching Everything (Phase 3)',
     'With standard depth, linearization maps all geometry to viewZ\u224815.0 (near plane). '
     'Far-plane threshold of 16.0 caught everything as "sky." Temporarily set to 0.0, '
     'then restored to 16.0 when reverting to reversed-Z.'),
]

for title, desc in issues:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

# ============================================================
# 6. CURRENT STATE
# ============================================================
doc.add_heading('6. Current State (2026-03-20)', level=1)

doc.add_heading('What Works', level=2)
works = [
    'Proxy: Full D3D11 wrapping, DLL loads correctly, no crashes',
    'Phase detection: Detects all 9 phases (but timing is wrong \u2014 see 4.3)',
    'Mid-frame dispatch: Fires callbacks, saves/restores state correctly',
    'HiZ pyramid: Compiles, builds mip chain, converts standard\u2192reversed-Z',
    'Linear depth: Compiles, builds from HiZ',
    'All 28 effect renderers: Compile successfully',
    'All 84 HLSL shaders: Compile via hot-reload',
    'Developer tools: Shader error overlay, live source viewer, GPU profiler, frame capture, debug visualizations',
    'Configuration: INI persistence, 5 user presets, hot-reload',
    'ImGui overlay: Full debug GUI',
    'SceneCompositor: Copy+blend pipeline works, RT format guard works',
]
for item in works:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('What Doesn\'t Work', level=2)
broken = [
    'No effect produces correct visible output. Every effect depends on depth, and depth is not reliably available at the correct time.',
    'Phase detector timing: Fires GeometryMain \u2192 PostProcess before all geometry is drawn',
    'Depth buffer access: No approach has reliably provided shader-readable depth at mid-frame time',
    'All temporal accumulation: Bypassed/disabled pending motion vector availability',
    'All spatial denoising: Bypassed pending correct depth for bilateral edge stopping',
    'SSGI voxelization: Theoretically functional but never tested with correct depth',
    '13 deferred-init renderers: Never initialized (Volumetric Clouds, Frame Gen, TSR, etc.)',
]
for item in broken:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 7. OPTIONS MOVING FORWARD
# ============================================================
doc.add_heading('7. Options Moving Forward', level=1)

options = [
    ('Option A: Fix Phase Detection (Recommended First Step)',
     [
         'Replace the Draw(3, 0) heuristic with a more reliable detection method:',
         '1. RT-change based detection: Track when the game unbinds the main depth DSV. '
         'Skyrim\'s PostProcess passes don\'t use depth. The transition from "main DSV bound" '
         'to "no DSV" is a reliable GeometryMain \u2192 PostProcess signal.',
         '2. Shader hash-based detection: Profile the game for a few hundred frames, record '
         'shader hashes at each transition point. Identify the specific shader hash that marks '
         'the end of geometry. This is how ENB does it.',
         '3. Draw count + RT heuristic combined: Instead of just counting draws, require BOTH '
         'draw count threshold AND an RT change (main DSV unbound) to trigger the transition.',
         'Effort: Medium. The RenderPhaseDetector is clean and modular \u2014 adding a new detection '
         'strategy is straightforward.',
     ]),
    ('Option B: Depth Copy at PostGeometry',
     [
         'If phase detection is fixed (effects fire after all geometry), then the depth at '
         'PostGeometry time is complete. The problem becomes purely about creating a '
         'shader-readable copy:',
         '1. CopyResource the depth texture to our R24G8_TYPELESS texture with BIND_SHADER_RESOURCE',
         '2. Create SRV with R24_UNORM_X8_TYPELESS format',
         '3. Build HiZ from this SRV',
         'Effort: Low. The copy infrastructure exists, just needs to fire at the right time.',
     ]),
    ('Option C: ReShade-Style Depth Interception',
     [
         'ReShade solves this by:',
         '1. Hooking CreateTexture2D to intercept the game\'s depth texture creation',
         '2. Modifying the bind flags to add BIND_SHADER_RESOURCE',
         '3. Creating the texture with R24G8_TYPELESS format instead of D24_UNORM_S8_UINT',
         'This gives zero-cost SRV access to the game\'s live depth buffer at any time \u2014 '
         'no copies, no ownership substitution.',
         'Effort: Medium. Requires modifying WrappedDevice::CreateTexture2D to detect the main '
         'depth texture (by size/format/bind flags) and silently upgrade its format and bind flags.',
     ]),
    ('Option D: Present-Time Only (Current \u2014 Verify)',
     [
         'If one-frame-old depth is acceptable, verify the current approach works. The last '
         'test was inconclusive. The issue may be: AcquireDepthSRV failing to find the bound '
         'DSV at Present time; the game\'s depth being cleared before Present; the SRV creation '
         'failing silently.',
         'Effort: Low (diagnostic only). Add more logging to AcquireDepthSRV and BuildPyramid.',
     ]),
    ('Option E: Nuclear \u2014 Full Depth Pipeline Rewrite',
     [
         'Start the depth pipeline from scratch using the ReShade approach (Option C) as the '
         'foundation:',
         '1. Intercept depth texture creation',
         '2. Build HiZ at PostGeometry (if phase detection is fixed) or Present',
         '3. Verify with diagnostic shader',
         '4. Then bring up effects one at a time',
         'Effort: High but definitive.',
     ]),
]

for title, paras in options:
    doc.add_heading(title, level=2)
    for para in paras:
        doc.add_paragraph(para)

# ============================================================
# 8. RECOMMENDED PATH
# ============================================================
doc.add_heading('8. Recommended Path', level=1)

steps = [
    'Fix phase detection (Option A) \u2014 this is the root cause of most issues',
    'Implement depth texture interception (Option C) \u2014 eliminates the depth access problem entirely',
    'Verify GTAO with diagnostic shader \u2014 confirm depth varies across the full screen at all camera angles',
    'Restore real GTAO shader \u2014 test actual AO output',
    'Bring up effects one at a time per the existing ROADMAP Phase 2',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}', style='List Number')

# ============================================================
# 9. CODE QUALITY NOTES
# ============================================================
doc.add_heading('9. Code Quality Notes', level=1)

doc.add_heading('Strengths', level=2)
strengths = [
    'Clean separation between proxy and plugin',
    'D3D11StateBackup is thorough (OM, RS, IA, VS, PS+CBs, CS)',
    'GPUResource.h eliminates boilerplate',
    'External HLSL with hot-reload is excellent for iteration',
    'RT format guard prevents compositor crashes',
    'SEH wrapper prevents total plugin crashes from AVs',
    'NaN/Inf sanitization on all game data',
]
for s in strengths:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Weaknesses', level=2)
weaknesses = [
    'Effects have shaders embedded as C++ string literals AND as external .hlsl files \u2014 dual maintenance burden',
    'Many renderers have identical patterns (CB upload, SRV acquisition, dispatch) that could be abstracted',
    'RenderPhaseDetector heuristics are fragile and game-version-specific',
    'No automated tests \u2014 everything is manual in-game testing',
    '13 renderers are deferred-init and have never been tested',
    'Temporal accumulation is disabled on ALL effects \u2014 motion vectors exist in code but aren\'t wired up',
]
for w in weaknesses:
    doc.add_paragraph(w, style='List Bullet')

doc.add_heading('Architectural Risks', level=2)
risks = [
    'The proxy wraps 100+ ID3D11DeviceContext methods \u2014 any missed method that modifies state can desync the state cache',
    'Phase change callbacks fire synchronously on the render thread \u2014 any slow effect will stall the game',
    'ProxyInterface is a flat C struct with raw function pointers \u2014 fragile ABI',
]
for r in risks:
    doc.add_paragraph(r, style='List Bullet')

# ============================================================
# 10. BUILD & DEPLOY
# ============================================================
doc.add_heading('10. Build & Deploy', level=1)

doc.add_heading('Build', level=2)
add_code_block("""cmake -B build -S . -DCMAKE_TOOLCHAIN_FILE=<vcpkg>/scripts/buildsystems/vcpkg.cmake -DVCPKG_TARGET_TRIPLET=x64-windows-static
cmake --build build --config Release""")

doc.add_heading('Deploy (MO2)', level=2)
add_code_block("""<mod>/SKSE/plugins/RAW.dll           -- SKSE plugin
<mod>/SKSE/plugins/RAW/Shaders/*.hlsl -- 84 hot-reloadable shaders
<mod>/SKSE/plugins/RAW/RAW.ini       -- persistent config
<mod>/Root/d3d11.dll                  -- proxy (MO2 ROOT mapping)""")

doc.add_heading('Dependencies', level=2)
doc.add_paragraph('CommonLibSSE-NG (colorglass registry)', style='List Bullet')
doc.add_paragraph('Dear ImGui (DX11 + Win32 bindings)', style='List Bullet')
doc.add_paragraph('VS2022, C++23, x64-windows-static', style='List Bullet')

# ============================================================
# 11. KEY FILES
# ============================================================
doc.add_heading('11. Key Files to Start With', level=1)

doc.add_paragraph('If you\'re picking this up, read in this order:')

key_files = [
    ('src/d3d11_proxy/proxy_main.cpp', 'How the proxy loads, wraps, and exposes its interface'),
    ('src/d3d11_proxy/RenderPhaseDetector.cpp', 'The phase detection heuristics (this is where the timing bug lives)'),
    ('src/core/PhaseDispatcher.cpp', 'How phase transitions trigger effect dispatch'),
    ('src/core/HiZPyramid.cpp', 'How depth is captured and converted'),
    ('src/core/GTAORenderer.cpp', 'Representative effect renderer (the first one to get working)'),
    ('src/core/SceneCompositor.cpp', 'How effect outputs are composited onto the game scene'),
]

for i, (path, desc) in enumerate(key_files, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {path}')
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.add_run(f' \u2014 {desc}')

# Footer
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'This document reflects the project state as of 2026-03-20. The codebase compiles and '
    'loads without crashes. The blocking issue is depth buffer access timing, which prevents '
    'all effects from producing correct output.'
)
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Save
output_path = r'C:\Users\Zain\SKSE\Playground\RAW\RAW_Technical_Handoff.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
